import os
from pdf2image import convert_from_path
import pytesseract
import spacy
from collections import Counter
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from PIL import Image
import fitz  # PyMuPDF
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

# Load French spaCy model
nlp = spacy.load("fr_core_news_sm")

# Chapter split pattern
CHAPTER_PATTERN = re.compile(r'(Chapitre\s+\d+)', re.IGNORECASE)

def split_text_into_chapters(text):
    splits = CHAPTER_PATTERN.split(text)
    chapters = []
    for i in range(1, len(splits), 2):
        header = splits[i].strip()
        body = splits[i+1].strip() if i+1 < len(splits) else ""
        chapters.append(f"{header}\n{body}")
    return chapters

def extract_words_spacy(text):
    # Normalize curly apostrophes and similar variants to straight apostrophe
    text = text.replace("’", "'").replace("‛", "'").replace("‘", "'")

    doc = nlp(text)
    words = []
    i = 0
    while i < len(doc):
        token = doc[i]
        # Combine contractions like d'eau, j'arrive, qu'on
        if token.text.endswith("'") and i + 1 < len(doc):
            combined = token.text + doc[i + 1].text
            words.append(combined.lower())
            i += 2
        else:
            if token.is_alpha:
                words.append(token.text.lower())
            i += 1
    return words

def save_docx_output(chapters_data, output_docx):
    """
    chapters_data = list of tuples: (chapter_number, [(word, count), ...])
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for chapter_num, words_counts in chapters_data:
        doc.add_heading(f'Chapter {chapter_num} Unique Words', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'  # Ensure table borders are visible
        table.autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'S.No'
        hdr_cells[1].text = 'French Word'
        hdr_cells[2].text = 'Occurrences'
        hdr_cells[3].text = 'English Translation'

        # Set column widths (in inches)
        widths = [Inches(0.7), Inches(1.5), Inches(1.1), Inches(3.0)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Helper to set cell padding and row height
        def set_cell_format(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Set cell padding
            for side in ('top', 'bottom', 'left', 'right'):
                node = OxmlElement(f'w:{side}Margin')
                node.set(qn('w:w'), '120')  # 120 twips = ~0.08 inch
                node.set(qn('w:type'), 'dxa')
                tcPr.append(node)

        for i, (word, count) in enumerate(words_counts, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = word
            row_cells[2].text = str(count)
            row_cells[3].text = ''  # blank for translation
            # Set row height (min height)
            tr = table.rows[-1]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '800')  # 800 twips = ~0.56 cm
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)
            # Set cell padding for all cells in the row
            for cell in row_cells:
                set_cell_format(cell)

        doc.add_paragraph()  # Space between chapters

    doc.save(output_docx)
    print(f"Saved output DOCX to '{output_docx}'")

def save_french_words_docx(chapters_data, output_docx):
    """
    chapters_data = list of tuples: (chapter_number, [(word, count), ...])
    Creates a DOCX with only the French words, one per line, grouped by chapter.
    """
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for chapter_num, words_counts in chapters_data:
        doc.add_heading(f'Chapter {chapter_num} French Words', level=1)
        for word, _ in words_counts:
            doc.add_paragraph(word)
        doc.add_paragraph()  # Space between chapters

    doc.save(output_docx)
    print(f"Saved French words DOCX to '{output_docx}'")

def save_french_words_excel(chapters_data, output_xlsx):
    """
    chapters_data = list of tuples: (chapter_number, [(word, count), ...])
    Creates an Excel file with French words, one per row, grouped by chapter.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "French Words"
    
    # Set up headers
    ws['A1'] = "Chapter"
    ws['B1'] = "French Word"
    ws['C1'] = "Occurrences"
    
    # Style the headers
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    
    # Add data
    row = 2
    for chapter_num, words_counts in chapters_data:
        for word, count in words_counts:
            ws[f'A{row}'] = f"Chapter {chapter_num}"
            ws[f'B{row}'] = word
            ws[f'C{row}'] = count
            row += 1
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
        ws.column_dimensions[column_letter].width = adjusted_width
    
    wb.save(output_xlsx)
    print(f"Saved French words Excel file to '{output_xlsx}'")

def get_images_from_folder(folder_path):
    exts = ('.jpg', '.jpeg', '.png', '.bmp', '.tiff')
    image_files = [f for f in os.listdir(folder_path) if f.lower().endswith(exts)]
    image_files.sort()  # Optional: sort by filename
    images = []
    for img_file in image_files:
        img_path = os.path.join(folder_path, img_file)
        try:
            img = Image.open(img_path)
            images.append(img)
        except Exception as e:
            print(f"Could not open {img_path}: {e}")
    return images

def extract_text_from_pdf(pdf_path):
    """
    Try to extract text directly from a PDF using PyMuPDF (fitz).
    Returns a list of strings, one per page.
    """
    text_pages = []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text = page.get_text()
            text_pages.append(text)
    return text_pages

def main():
    print("Choose input type:")
    print("1. PDF file")
    print("2. Folder of images")
    input_type = input("Enter 1 for PDF or 2 for images: ").strip()

    if input_type == '1':
        input_pdf = input("Enter full path to input PDF file (e.g. C:\\Users\\You\\Desktop\\book.pdf): ").strip()
        if not os.path.isfile(input_pdf):
            print("Error: Input PDF file does not exist. Exiting.")
            return
        try:
            print("Trying to extract text directly from PDF...")
            text_pages = extract_text_from_pdf(input_pdf)
            # Check if most pages have enough text (not just whitespace)
            if sum(len(t.strip()) > 30 for t in text_pages) > 0.7 * len(text_pages):
                print("PDF appears to be text-based. Using direct text extraction.")
                pages = text_pages
                is_text_based = True
            else:
                print("PDF appears to be image-based. Using OCR.")
                pages = convert_from_path(input_pdf, dpi=600)
                is_text_based = False
        except Exception as e:
            print(f"Error during PDF text extraction: {e}\nFalling back to OCR.")
            pages = convert_from_path(input_pdf, dpi=600)
            is_text_based = False
    elif input_type == '2':
        input_folder = input("Enter full path to folder containing images: ").strip()
        if not os.path.isdir(input_folder):
            print("Error: Input folder does not exist. Exiting.")
            return
        pages = get_images_from_folder(input_folder)
        if not pages:
            print("No images found in the folder. Exiting.")
            return
        print(f"Found {len(pages)} images.")
        is_text_based = False
    else:
        print("Invalid input type. Exiting.")
        return

    output_folder = input("Enter full path to output folder where result XLSX will be saved: ").strip()
    if not os.path.isdir(output_folder):
        print("Error: Output folder does not exist. Exiting.")
        return

    print(f"Processing {len(pages)} pages/images...")
    full_text = ""
    for i, page in enumerate(pages, 1):
        if is_text_based:
            text = page
        else:
            text = pytesseract.image_to_string(page, lang="fra")
        full_text += text + "\n"

    print("Splitting text into chapters...")
    chapters = split_text_into_chapters(full_text)
    if not chapters:
        print("Warning: No chapters found! Saving entire text as one chapter.")
        chapters = [full_text]

    print(f"Found {len(chapters)} chapters.")

    chapters_data = []
    for idx, chapter_text in enumerate(chapters, 1):
        print(f"Processing Chapter {idx}...")
        words = extract_words_spacy(chapter_text)
        word_counts = Counter(words)
        unique_words = sorted(word_counts.items())  # list of (word, count)
        chapters_data.append((idx, unique_words))

    # Save French words to Excel
    output_french_excel = os.path.join(output_folder, "output_french_words_only.xlsx")
    save_french_words_excel(chapters_data, output_french_excel)

    print("Done! Check the output Excel file for results.")

if __name__ == "__main__":
    main()
